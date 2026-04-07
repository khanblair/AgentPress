"""
app/tools/document_skills/docx_builder.py

Builds a brand-compliant .docx file from structured draft text.
Parses "## Section" markers from draft_text.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRAND_HEADING_HEX = os.environ.get("BRAND_PRIMARY_HEX", "1A1A2E")
BRAND_ACCENT_HEX = os.environ.get("BRAND_ACCENT_HEX", "E94560")
BRAND_BODY_HEX = os.environ.get("BRAND_TEXT_HEX", "1A1A1A")
BRAND_FONT = os.environ.get("BRAND_FONT", "Calibri")


def _hex_to_rgb(hex_str: str) -> RGBColor:
    hex_str = hex_str.lstrip("#")
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _parse_sections(draft_text: str) -> list[dict]:
    pattern = re.compile(r"^##\s+(.+)$", re.MULTILINE)
    sections = []
    matches = list(pattern.finditer(draft_text))
    for i, match in enumerate(matches):
        heading = match.group(1).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(draft_text)
        body = draft_text[start:end].strip()
        sections.append({"heading": heading, "body": body})
    if not sections:
        sections = [{"heading": "Document", "body": draft_text}]
    return sections


def build_docx(draft_text: str, output_path: str) -> bool:
    """Build a brand-compliant .docx. Returns True on success."""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title page heading
    title_para = doc.add_heading("AgentPress Document", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.name = BRAND_FONT
    title_run.font.color.rgb = _hex_to_rgb(BRAND_HEADING_HEX)
    title_run.font.size = Pt(28)

    doc.add_page_break()

    sections = _parse_sections(draft_text)
    for sec in sections:
        # Section heading
        h = doc.add_heading(sec["heading"], level=1)
        h_run = h.runs[0] if h.runs else h.add_run(sec["heading"])
        h_run.font.name = BRAND_FONT
        h_run.font.color.rgb = _hex_to_rgb(BRAND_ACCENT_HEX)
        h_run.font.size = Pt(16)
        h_run.bold = True

        # Body paragraphs
        for line in sec["body"].splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith(("•", "-", "*")):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line.lstrip("•-* ").strip())
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
            run.font.name = BRAND_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = _hex_to_rgb(BRAND_BODY_HEX)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    assert Path(output_path).exists(), f"docx_builder: Output not found at {output_path}"
    return True
